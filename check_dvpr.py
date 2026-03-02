#!/usr/bin/env python3
"""Check DVPR template and whether deliverables require it."""
import sys
sys.path.insert(0, '/home/chu2026/Documents/github/APQP/.venv/lib/python3.13/site-packages')

import openpyxl

SPEC = "/home/chu2026/Documents/github/APQP/Spec/FUEL BUNDLE AND FUEL SYSTEM"

# 1. DVPR template
print("=" * 60)
print("1. STLA-DVPR模板.xlsx")
print("=" * 60)
wb = openpyxl.load_workbook(f"{SPEC}/STLA-DVPR模板.xlsx")
for sn in wb.sheetnames:
    ws = wb[sn]
    print(f"\n--- Sheet: {sn} (rows={ws.max_row}, cols={ws.max_column}) ---")
    for row in ws.iter_rows(min_row=1, max_row=min(15, ws.max_row), values_only=True):
        vals = [str(c)[:80] if c else '' for c in row]
        if any(v for v in vals):
            print(vals)

# 2. Deliverables checklist
print("\n" + "=" * 60)
print("2. KP1_Deliverables - searching for DVPR/DVP&R/test plan")
print("=" * 60)
wb2 = openpyxl.load_workbook(f"{SPEC}/KP1_Deliverables_Supplier-Checklist.xlsx")
for sn in wb2.sheetnames:
    ws = wb2[sn]
    print(f"\n--- Sheet: {sn} ---")
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
        vals = [str(c) if c else '' for c in row]
        line = ' | '.join(vals)
        print(line)
        for kw in ['DVPR', 'DVP', 'DV plan', 'PV plan', 'validation', 'test plan', 'DVPROBAN']:
            if kw.lower() in line.lower():
                print(f"  >>> FOUND: '{kw}'")

# 3. TDR docx
print("\n" + "=" * 60)
print("3. KP1_TDR - searching for DVPR/DVP&R/test plan")
print("=" * 60)
try:
    from docx import Document
    doc = Document(f"{SPEC}/KP1_TDR_Technical-Offer-Documents.docx")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
            for kw in ['DVPR', 'DVP', 'DV plan', 'PV plan', 'validation', 'test plan']:
                if kw.lower() in p.text.lower():
                    print(f"  >>> FOUND: '{kw}'")
    for t in doc.tables:
        for row in t.rows:
            vals = [c.text.strip() for c in row.cells]
            line = ' | '.join(vals)
            print(line)
            for kw in ['DVPR', 'DVP', 'DV plan', 'PV plan', 'validation', 'test plan']:
                if kw.lower() in line.lower():
                    print(f"  >>> FOUND: '{kw}'")
except Exception as e:
    print(f"Error: {e}")
